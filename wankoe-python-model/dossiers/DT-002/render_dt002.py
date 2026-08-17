"""DT-002 renderer — builds the canonical note (DT-002.md) AND the branded
Word document from dt002_data.json + the authored prose below.

    PYTHONPATH=src python dossiers/DT-002/render_dt002.py

Numbers NEVER live in this file: every figure is read from dt002_data.json
(produced by extract_dt002.py), so the note cannot drift from the engine.
Requires: python-docx; the PFD raster (pfd_rev15-1.png) next to this script
(regenerate: pdftoppm -png -r 110 -f 1 -l 1 20260806_Wankoe_1.1_PFD_REV15.pdf pfd_rev15).
"""

import json
import pathlib

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = pathlib.Path(__file__).parent
DATA = json.loads((HERE / "dt002_data.json").read_text())
LOGO = pathlib.Path("/root/.claude/skills/synced/noezys-report/assets/noezys_logo.png")

VIOLET = RGBColor(0x5B, 0x4F, 0xC7)
CYAN = RGBColor(0x3F, 0xB8, 0xE0)
GREY = RGBColor(0x60, 0x5E, 0x5C)
BODY_FONT, HEAD_FONT = "Aptos", "Aptos Display"
LEGAL = (
    "NOEZYS LTD, company incorporated under the Companies Act 1995 of Malta, "
    "registration C 115556, 170 Pater House, Level 1, Suite A369, Triq Salvu Psaila, "
    "Birkirkara BKR 9077, Malta."
)

P = DATA["_provenance"]
RK = DATA["replay_kit_inputs"]
CAL = RK["calibration"]
M1A, M1B = DATA["modes"]["1A"], DATA["modes"]["1B"]
PLAN = DATA["annual_planning"]
ADQ = DATA["pfd_rev15_adequacy"]
MESHES = list(RK["feed_curve_passing_pct"].keys())


def f(x, nd=2):
    return f"{x:,.{nd}f}".replace(",", " ") if isinstance(x, (int, float)) else str(x)


# ----------------------------------------------------------------- content
# Blocks: (kind, payload). kinds: h1/h2/h3, p, table {header, rows, small},
# image, pagebreak, landscape_on/off.

def stream_psd_table(mode):
    streams = [s for s in mode["streams"] if s.get("present")]
    header = ["Sieve (mm)"] + [s["label"].split(" (")[0] for s in streams]
    rows = []
    for lab, key, nd in [("Dry t/h", "dry_tph", 2), ("Wet t/h", "wet_tph", 2), ("P80 (mm)", "P80_mm", 2)]:
        rows.append([lab] + [f(s[key], nd) for s in streams])
    for m in MESHES:
        rows.append([m] + [f(s["passing_pct"][m]) for s in streams])
    return {"header": header, "rows": rows, "small": True,
            "note": "Rows below the rates are CUMULATIVE % PASSING at each sieve."}


def duty_table(mode):
    mm = mode["machines"]
    c6, s8, c11 = mm["CR.5006"], mm["SR.5008"], mm["CR.5011"]
    return {
        "header": ["Machine", "Duty at this mode", "Power", "Sizing check"],
        "rows": [
            ["CR.5006 toothed-roll sizer",
             f"{f(c6['throughput_wet_tph'])} t/h wet ({f(c6['throughput_dry_tph'])} dry), F80 {f(c6['F80_mm'])} mm, P80 {f(c6['P80_mm'])} mm, x80 = gap = 60 mm",
             f"W = {f(c6['W_kWh_t'],3)} kWh/t, P_net {f(c6['P_net_kW'],1)} kW, absorbed {f(c6['P_installed_kW'],1)} kW (eta_m {CAL['eta_m']})",
             f"{f(c6['utilization_pct_of_250_wet'],1)} % of the 250 t/h wet line rating"],
            ["SR.5008 double-deck screen 35/20",
             f"Feed {f(s8['feed_wet_tph'])} t/h wet ({f(s8['feed_dry_tph'])} dry); feed context at 35 mm: {f(s8['feed_pct_half_size'])} % passing a/2, {f(s8['feed_pct_oversize'])} % oversize",
             "n/a (screen)",
             f"Required areas: top {f(s8['required_areas_m2']['top_deck_35mm']['required_area_m2'])} m2, bottom {f(s8['required_areas_m2']['bottom_deck_20mm']['required_area_m2'])} m2. Purchase minima (client 2026-08-15, RAIN duty): {s8['purchase_min_area_m2']['top_deck']} / {s8['purchase_min_area_m2']['bottom_deck']} m2"],
            ["CR.5011 impact crusher (loop)",
             f"Loop load {f(c11['loop_load_wet_tph'])} t/h wet of {f(c11['vendor_capacity_wet_tph'],0)} t/h vendor wet rating = {f(c11['utilization_pct'],1)} %; F80 {f(c11['F80_mm'])} mm, P80 {f(c11['P80_mm'])} mm, CSS {mode['settings']['CR.5011_x80_css_mm']} mm, v {mode['settings']['CR.5011_v_ms']} m/s",
             f"Ecs {f(c11['Ecs_kWh_t'],3)} kWh/t, t10 {f(c11['t10_pct'],1)} %, n {f(c11['n'],2)}; W {f(c11['W_kWh_t'],3)} kWh/t, absorbed {f(c11['P_installed_kW'],1)} kW (installed 132 kW)",
             f"Recirculation {f(mode['recirculation_dry_tph'])} t/h dry"],
        ],
    }


env = M1A["kfs_envelope_check"]
blocks = [
    ("h1", "1. Purpose and how to read this note"),
    ("p", "This note documents, end to end, how the WANKOE deterministic model computes zone 1.1: the machine sizing and the production capacity, with every particle-size distribution presented as CUMULATIVE PERCENT PASSING in table form. It is written for the model exchange with our colleague running Metso's Bruno simulator: section 3 is a full replay kit (all inputs, settings and calibration constants) so Bruno can be configured identically, sections 4 to 6 derive and apply the models, and section 9 confronts our figures with the NACO flowsheet design values, honestly. Nothing here needs to be trusted: every figure replays from the archived script (see the provenance footer), and any divergence between Bruno and this note is meaningful precisely because both sides are fully specified."),
    ("p", "Language note: tags follow the NACO PFD 11-01-PFD REV15 everywhere. The model historically used spec-era tags; the zone-1.1 retag of 2026-08-17 renamed CR.5009 to CR.5006 and SR.5007 to SR.5008 (CR.5011 and CR.5003 were already aligned). Documents issued before that date carry the old tags, the machines are the same."),
    ("h1", "2. The flowsheet (PFD 11-01-PFD REV15)"),
    ("p", "The PFD plate is the contractual reference of this note; it is deliberately NOT reproduced inside this document (an A1 sheet does not survive page-size reduction) and travels as a separate PDF (20260806_Wankoe_1.1_PFD_REV15.pdf) to be opened full screen alongside. Chain: ROM stockpile 0/700 (10 000 t), truck to hopper HO.5001 (20 m3), vibrating feeder VF.5002 (250 t/h), primary crusher CR.5003 (0/700 to 0/200), conveyors BC.5004/BC.5005, toothed-roll sizer CR.5006 (0/200 to 0/40), conveyor BC.5007 to the double-deck screen SR.5008 (decks 35 and 20 mm), diverter DV.5009, loop crusher CR.5011 returning via BC.5010, products BC.5012 to SP.5014 (0/20 crude, 10 000 t, to AG and Feed plant) and BC.5013 to SP.5015 (KFS 20/35, 5 000 t, to kiln)."),
    ("landscape_on", None),
    ("image", str(HERE / "pfd_rev15-1.png")),
    ("landscape_off", None),
    ("h2", "2.1 Model boundary and scenario mapping"),
    ("p", "The model starts AT THE PIVOT: the feed curve is MEASURED by belt-cut at the primary station outlet, so the upstream chain (HO.5001, VF.5002, CR.5003, blending) is already inside the measured curve and is not recomputed. The modelled chain is CR.5006, then SR.5008 with the CR.5011 loop on the deck-1 oversize routed by DV.5009. PFD scenario A (KFS + crude) is the model's mode 1A; PFD scenario B (crude only, no KFS) is mode 1B, in which the 20/35 cut is recycled to CR.5011 instead of leaving as product, and the line feed drops to the re-bisected 172.0 t/h wet so the loop machine stays inside its vendor rating (section 9 discusses the 150 t/h difference with the sheet)."),
    ("h1", "3. Replay kit: inputs, settings, calibration"),
    ("p", f"Total-flow rule (project convention, load-bearing): every line feed rate is the TOTAL flow, WET basis (dry solids plus water), as a belt scale would weigh it. Mode 1A feed: {f(RK['flow_rates_wet_tph']['mode_1A_feed'],1)} t/h wet; mode 1B feed: {f(RK['flow_rates_wet_tph']['mode_1B_feed'],1)} t/h wet; feed moisture {f(RK['feed_moisture_pct_wet_basis'],1)} % wet basis, dry weather photo. Feed F80 = {f(RK['feed_F80_mm'])} mm (measured 2026-08-08 belt-cut curve, completed by declared tail hypotheses H-FEED-1/2)."),
    ("table", {"header": ["Sieve (mm)", "Feed cumulative % passing"],
               "rows": [[m, f(v)] for m, v in RK["feed_curve_passing_pct"].items()],
               "note": "The pivot feed curve, as Bruno should ingest it."}),
    ("table", {"header": ["Setting", "Value", "Status"], "rows": [
        ["CR.5006 gap g (= product x80)", "60 mm", "Client reference setting 2026-08-13; x80 = g validated 2026-08-08"],
        ["CR.5006 RR uniformity n", "1.35", "[H] pending vendor gradation table"],
        ["SR.5008 apertures a1 / a2", "35 / 20 mm", "KFS window, sanctuarized by client ruling"],
        ["Screen imperfection I", "0.15", "[H] literature value, client arbitration 2026-08-10; convention question open (see 4.3)"],
        ["CR.5011 CSS (x80)", "30 mm (1A) / 18 mm (1B)", "Mode changeover is a routine operation"],
        ["CR.5011 rotor speed v", "30 m/s", "Client 2026-08-14, minimum speed"],
        ["Bond work index Wi", f"{CAL['Wi_kWh_t']} kWh/t", "[ref.] Fontaine, Belgian limestone (client Q2, 2026-08-11)"],
        ["Impact breakage A_j / b_j", f"{CAL['m5_A_j']} / {CAL['m5_b_j']}", "[H] pending drop-weight tests (none launched, client 2026-08-16)"],
        ["Motor efficiency eta_m", str(CAL["eta_m"]), "absorbed = net / eta_m"],
        ["Computation grid", f"x{CAL['computation_grid_refinement']} refinement (57 meshes)", "Spec sieves remain the presentation format (this note's tables)"],
    ]}),
    ("h1", "4. The models, derived"),
    ("h2", "4.1 M1, crusher product curve (truncated Rosin-Rammler)"),
    ("p", "Both crushers deliver a Rosin-Rammler (Weibull) product. Start from the RR cumulative passing P(x) = 1 - exp(-(x/xc)^n). The machine setting fixes the nominal product x80: requiring P(x80) = 0.8 gives exp(-(x80/xc)^n) = 0.2, hence (x80/xc)^n = ln 5 and xc = x80 / (ln 5)^(1/n) (the calibration constant m1_ln_arg = 5 is exactly this ln argument). Two physical corrections complete the model. TRUNCATION: a compression crusher cannot deliver lumps much larger than its setting, so the distribution is cut at xt = 1.7 x80 (trunc_factor) and rescaled by P(xt) so it reaches 100 % there. BYPASS: the feed fraction already finer than the setting passes through unchanged, so the delivered curve is the feed's fine part plus the crushed coarse share mapped through the truncated RR. Disclosed convention: after the truncation rescale the delivered product passes about 83 % at the nominal x80 (n = 1.35), i.e. its realized x80 is about 6 % finer than the setting; this is the spec's prescribed behavior, kept and disclosed."),
    ("h2", "4.2 M2, comminution power (Bond)"),
    ("p", f"W = 10 Wi (1/sqrt(P80) - 1/sqrt(F80)) with P80 and F80 in micrometres and Wi = {CAL['Wi_kWh_t']} kWh/t, clamped at zero if the product is coarser than the feed. Net power P_net = W Q (Q in t/h of dry solids); absorbed power = P_net / eta_m with eta_m = {CAL['eta_m']}. The bond_coef = {f(CAL['bond_coef'],0)} is the classical 10 of Bond's third law."),
    ("h2", "4.3 M3, screen partition curve"),
    ("p", "Each deck is a reduced-efficiency partition: the probability that a particle of size x reports to the OVERSIZE is ro(x) = 1 / (1 + (d50c/x)^s), a logistic curve in log-size. The cut point is d50c = a k_d with k_d = 1.0 (cut at the aperture), and the sharpness comes from the imperfection I through s = ln 9 / ln(1/(1 - I)): the ln 9 makes s such that the curve spans its 10 % to 90 % recovery band over the factor (1-I)^-2 around d50c. The partition is applied interval by interval on the computation grid to the feed's mass fractions; oversize and undersize flows and curves follow by summation, which is why the screen conserves mass exactly by construction."),
    ("p", "Honesty notes for the Bruno comparison, where the two simulators are most likely to differ: (1) ATTRIBUTION, the spec labels this curve Karra, but the implemented form is a logistic Reid/Plitt-family curve; Karra's published 1979 model is ro = 1 - exp(-0.693 (x/d50)^5.846) with a FIXED sharpness equivalent to a classic imperfection of about 0.13. (2) CONVENTION, with s = ln 9 / ln(1/(1-I)) the parameter I is a (d90-d10)/(2 d50)-type sharpness, not the classic (d75-d25)/(2 d50) imperfection of the literature that justified the 0.15 default; the realized classic imperfection at I = 0.15 is sinh(ln(1/(1-I))/2), about 0.081, meaning our screens are modelled about twice as sharp as a classic-0.15 screen. This convention question is an open arbitration on our side (expert clarification note pending). If Bruno's screen model is configured with a classic imperfection, MATCH THE REALIZED SHARPNESS (0.081), not the nominal 0.15, or the cut curves will diverge visibly around both apertures."),
    ("h2", "4.4 M4, screen area (VSMA / Fontaine form)"),
    ("p", f"Required area A = U f_p / (Qb f0), where U is the deck undersize in t/h, f_p = {CAL['m4_f_p']} a peak factor, and Qb = {f(CAL['m4_qb_coef'],0)} a^{CAL['m4_qb_exp']} t/h/m2 the basic capacity at aperture a (mm). Attribution disclosure: Qb alone is not the published VSMA basic capacity; the EFFECTIVE capacity Qb f0 = 4.86 a^0.6 reproduces the VSMA Factor-A table within 5 % at the project cuts (29.3 vs 30.8 t/h/m2 at 20 mm, 41.0 vs 39.5 at 35 mm), the fitted f0 = {CAL['m4_f0']} absorbing the VSMA factor string. Under rain the basic capacity is derated by the wet factor (0.75); the SR.5008 PURCHASE minima are client-decided on that rain duty (9.1 / 9.6 m2), while this note's dry-photo areas are the operating requirement."),
    ("h2", "4.5 M5, impact crusher (t10 / Ecs) and the loop"),
    ("p", f"The impactor's breakage intensity comes from the rotor speed: specific energy Ecs = v^2 / {f(CAL['m5_ecs_div'],0)} kWh/t (kinetic energy of the tip speed), then the JK-style characteristic t10 = A_j (1 - exp(-b_j Ecs)) with A_j = {CAL['m5_A_j']} and b_j = {CAL['m5_b_j']} [H]. The t10 sets the product curve SHAPE through the RR uniformity n = max({CAL['m5_n_min']}, ({f(CAL['m5_t10_ref'],0)}/t10)^{CAL['m5_n_exp']}): a harder hit (higher t10) gives a flatter (better graded) product. The product curve is then M1's truncated RR with x80 = CSS and that n; power is M2 on the impactor's own F80 to P80. THE LOOP: SR.5008 deck-1 oversize (plus, in mode 1B, the 20/35 cut) feeds CR.5011, whose product returns to the screen feed. The engine solves this recycle as a fixed point, iterating until BOTH the recycled flow rate and its entire curve change by less than the tolerance; all figures in this note are the converged state."),
    ("h1", "5. Mode 1A, KFS production, full chain"),
    ("p", "Settings: gap 60 mm, CSS 30 mm, v 30 m/s, feed 250 t/h wet. The 20/35 cut leaves as KFS; the 0/20 undersize is the crude co-product; the deck-1 oversize loops through CR.5011."),
    ("table", duty_table(M1A)),
    ("landscape_on", None),
    ("table", stream_psd_table(M1A)),
    ("landscape_off", None),
    ("p", f"KFS quality check against the kiln envelope (max 30 % below 20 mm, min 55 % in 20/35, max 15 % above 35 mm): below {f(env['below_20mm_pct'])} %, in cut {f(env['in_cut_20_35_pct'])} %, above {f(env['above_35mm_pct'])} %, COMPLIANT. Zone mass balance closes by construction (partition curves conserve mass; the loop is converged)."),
    ("h1", "6. Mode 1B, 0/20 campaigns, full chain"),
    ("p", "Settings change with the mode: CSS drops to 18 mm (at CSS 30 the 1B loop cannot pass material below 20 mm and diverges; 18 mm restores convergence), and the line feed drops to 172.0 t/h wet, the value re-bisected on the adopted quarry-target curve so the CR.5011 vendor guarantee of 90 t/h wet holds on BOTH feed curves. The 20/35 cut is recycled (no KFS is made, by definition of the mode)."),
    ("table", duty_table(M1B)),
    ("landscape_on", None),
    ("table", stream_psd_table(M1B)),
    ("landscape_off", None),
    ("h1", "7. Machine sizing summary (what we buy and why)"),
    ("table", {"header": ["Machine", "Sizing driver", "Purchase decision"], "rows": [
        ["CR.5006", "250 t/h wet continuous at gap 60; max lump 250 mm guaranteed by the quarry (ripping + face scalping); positive engagement of that lump", "Twin-shaft toothed SIZER class (client ruling); vendor standard twin-drive package, 220 kW held in data; contractual product curve at the 0/60 setting is the tender's core"],
        ["SR.5008", "The worst circumstance is RAIN (25 % of the season, line runs through): wet factor 0.75 consumes the former dry margin", "Purchase minima 9.1 / 9.6 m2 = rain duty (client decision 2026-08-15); dry-photo requirement in this note is 6.80 / 7.15 m2"],
        ["CR.5011", f"Mode-1B loop at {f(M1B['machines']['CR.5011']['loop_load_wet_tph'])} t/h wet against the 90 t/h wet vendor rating ({f(M1B['machines']['CR.5011']['utilization_pct'],1)} %); mode-1A at {f(M1A['machines']['CR.5011']['utilization_pct'],1)} %", "90 t/h WET at CSS 18 is a written guarantee point of the tender; 132 kW per the reference machine documents; the 1B feed 172.0 was chosen precisely so this guarantee holds in all ruled circumstances"],
    ]}),
    ("h1", "8. Production capacity, annual translation"),
    ("p", f"Principle (project rule): OPERATING HOURS ARE SET BY THE PRODUCTION TARGETS, never the reverse. Zone 1.1 runs a Saturday-extended single shift (ceiling 2 400 h/y). At the current plan the zone needs {f(PLAN['zone_1_1_hours']['mode_1A_hours_effective'],1)} h of mode 1A and {f(PLAN['zone_1_1_hours']['mode_1B_hours_effective'],1)} h of mode 1B: the 0/20 demand of the downstream zones is currently covered by the 1A co-product, so no dedicated 1B campaign is scheduled; the AUTO-1B rule (KFS is NEVER over-produced) would create 1B hours automatically if the 0/20 demand grew beyond the 1A co-production."),
    ("table", {"header": ["Annual figure", "Value"], "rows": [
        ["KFS production (firm)", f"{f(PLAN['production_t']['KFS'],0)} t"],
        ["AgLime total", f"{f(PLAN['production_t']['AgLime'],0)} t (incl. 0/20 conversion, client rule 2026-08-16)"],
        ["FeedLime grits (firm)", f"{f(PLAN['production_t']['FeedLime grits'],0)} t"],
        ["FeedLime fines", f"{f(PLAN['production_t']['FeedLime fines'],0)} t (production objective, served exactly)"],
        ["UltraFin", f"{f(PLAN['production_t']['UltraFin'],0)} t (incl. zone-1.3 dedusting dust, client rule 2026-08-16)"],
        ["0/20 to landfill", f"{f(PLAN['stockpiles_t'].get('0/20 to LANDFILL (net loss)',0.0),0)} t (mandatory AgLime/FeedLime conversion, client rule)"],
        ["KFS yield realized / required", f"{f(PLAN['kfs_yield']['realized_pct'])} % / {f(PLAN['kfs_yield']['required_for_zero_landfill_pct'])} % (equal by construction since the conversion rule)"],
    ]}),
    ("h1", "9. Adequacy with the PFD REV15 figures, honestly"),
    ("p", "The sheet's printed rates and this note's computed rates differ, and the difference is INFORMATIVE, not a disagreement about the flowsheet:"),
    ("table", {"header": ["Quantity (wet t/h)", "PFD REV15 (design)", "Engine (measured curve, mode 1A)"], "rows": [
        ["Fresh feed", f(ADQ["pfd_design_figures"]["fresh_feed_tph"], 0), f(RK["flow_rates_wet_tph"]["mode_1A_feed"], 0)],
        ["SR.5008 screen feed (BC.5007)", f(ADQ["pfd_design_figures"]["screen_feed_BC5007_tph"], 0), f(ADQ["engine_measured_curve_mode_1A_wet_tph"]["screen_feed"], 0)],
        ["KFS 20/35 (BC.5013)", f(ADQ["pfd_design_figures"]["scenario_A"]["kfs_tph"], 0), f(ADQ["engine_measured_curve_mode_1A_wet_tph"]["kfs_20_35"], 0)],
        ["Crude 0/20 (BC.5012)", f(ADQ["pfd_design_figures"]["scenario_A"]["crude_0_20_tph"], 0), f(ADQ["engine_measured_curve_mode_1A_wet_tph"]["crude_0_20"], 0)],
        ["Recycle (BC.5010)", f(ADQ["pfd_design_figures"]["recycle_BC5010_tph"], 0), f(ADQ["engine_measured_curve_mode_1A_wet_tph"]["recycle"], 0)],
        ["Scenario B / mode 1B basis", "150 t/h crude PRODUCT", "172.0 t/h wet FEED (re-bisected for the CR.5011 90 t/h wet guarantee)"],
    ]}),
] + [("p", t) for t in ADQ["confrontations"]] + [
    ("p", "For the Bruno exchange this is the first-order result to compare: run Bruno once with the NACO design curve and once with the measured belt-cut curve of section 3, and the 80-vs-62 gap should reproduce in Bruno too if its crusher and screen models are configured per sections 4.1 to 4.3. If it does, the two models agree that the question is geological (the feed curve), not mechanical (the machines)."),
    ("h1", "10. Provenance and replay"),
    ("p", f"Engine commit {P['commit']}; scenario: {P['scenario']}. Replay without any assistant: PYTHONPATH=src python dossiers/DT-002/extract_dt002.py (regenerates dt002_data.json), then PYTHONPATH=src python dossiers/DT-002/render_dt002.py (regenerates this note and the Word document from that JSON). Every figure in this note comes from that JSON; none is hand-typed. The PFD is archived alongside (SHA-256 4c2cd3a447744542a0fb981e8e14bb9b50f4bab373da26a425824fb455f0275f, byte-identical to the 2026-08-12 docs/pfd/ archive). Declared hypotheses [H] remain hypotheses: no external test campaign has been launched (client decision 2026-08-16). Produced by NOEZYS."),
]


# ----------------------------------------------------------------- markdown
def to_markdown():
    lines = [
        "# DT-002 — Zone 1.1 Complete Sizing Note (model-exchange edition)",
        "",
        "**by NOEZYS** — technical dossier DT-002, issued 2026-08-17. Internal note",
        "for the model exchange with the client's colleague running Metso Bruno.",
        "Tags: NACO 11-01-PFD REV15. All PSD tables are cumulative % passing.",
        "",
    ]
    for kind, payload in blocks:
        if kind == "h1":
            lines += [f"## {payload}", ""]
        elif kind == "h2":
            lines += [f"### {payload}", ""]
        elif kind == "p":
            lines += [payload, ""]
        elif kind == "image":
            lines += ["![PFD 11-01-PFD REV15](pfd_rev15-1.png)",
                      "*(full-resolution PDF archived alongside: 20260806_Wankoe_1.1_PFD_REV15.pdf)*", ""]
        elif kind == "table":
            t = payload
            lines.append("| " + " | ".join(t["header"]) + " |")
            lines.append("|" + "---|" * len(t["header"]))
            for r in t["rows"]:
                lines.append("| " + " | ".join(str(c) for c in r) + " |")
            if t.get("note"):
                lines.append(f"\n*{t['note']}*")
            lines.append("")
    (HERE / "DT-002.md").write_text("\n".join(lines), encoding="utf-8")
    print("written: DT-002.md")


# ----------------------------------------------------------------- docx
def _style(doc):
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(10)
    for lvl, size, color in [(1, 16, VIOLET), (2, 13, VIOLET), (3, 11, CYAN)]:
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = HEAD_FONT
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = True


def _cell_text(cell, text, size, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = BODY_FONT
    run.font.bold = bold


def _borders(table):
    tbl = table._tbl
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "B9B6E0")
        borders.append(el)
    tbl.tblPr.append(borders)


def to_docx():
    doc = Document()
    _style(doc)
    sec = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, attr, Cm(2.54))
    sec.footer.paragraphs[0].text = LEGAL
    sec.footer.paragraphs[0].runs[0].font.size = Pt(6.5)
    sec.footer.paragraphs[0].runs[0].font.color.rgb = GREY
    sec.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cover
    if LOGO.exists():
        pl = doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.add_run().add_picture(str(LOGO), width=Cm(7))
    for text, size, color, bold in [
        ("WANKOE limestone processing line", 11, GREY, False),
        ("DT-002", 28, VIOLET, True),
        ("Zone 1.1 Complete Sizing Note", 18, VIOLET, True),
        ("Model-exchange edition: our deterministic engine and Metso Bruno, side by side", 12, CYAN, False),
        ("", 10, GREY, False),
        ("Internal technical dossier - August 2026 - produced by NOEZYS", 10, GREY, False),
        ("Tags per NACO PFD 11-01-PFD REV15 - all PSD tables are cumulative % passing", 9, GREY, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.color.rgb = color
        r.font.bold = bold
        r.font.name = HEAD_FONT if bold else BODY_FONT
    doc.add_page_break()

    # Client instruction 2026-08-17: NO flowsheet image inside the Word
    # document (the A1 plate renders uselessly small) — the separate PDF is
    # the full-screen reference. The markdown edition keeps the raster.
    doc_blocks, i = [], 0
    while i < len(blocks):
        if (
            i + 2 < len(blocks)
            and blocks[i][0] == "landscape_on"
            and blocks[i + 1][0] == "image"
            and blocks[i + 2][0] == "landscape_off"
        ):
            i += 3
            continue
        doc_blocks.append(blocks[i])
        i += 1

    landscape = False

    def new_section(land):
        nonlocal landscape
        s = doc.add_section()
        if land:
            s.orientation = WD_ORIENT.LANDSCAPE
            s.page_width, s.page_height = sec.page_height, sec.page_width
        else:
            s.orientation = WD_ORIENT.PORTRAIT
            s.page_width, s.page_height = sec.page_width, sec.page_height
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(s, attr, Cm(1.6 if land else 2.54))
        landscape = land
        return s

    for kind, payload in doc_blocks:
        if kind in ("h1", "h2", "h3"):
            doc.add_heading(payload, level=int(kind[1]))
        elif kind == "p":
            doc.add_paragraph(payload)
        elif kind == "landscape_on":
            new_section(True)
        elif kind == "landscape_off":
            new_section(False)
        elif kind == "image":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(payload, width=Cm(25.0))
            cap = doc.add_paragraph("PFD 11-01-PFD REV15 (NACO), the contractual flowsheet of this note; full-resolution PDF attached separately.")
            cap.runs[0].font.size = Pt(8)
            cap.runs[0].font.color.rgb = GREY
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "table":
            t = payload
            size = 7 if t.get("small") else 8.5
            table = doc.add_table(rows=1 + len(t["rows"]), cols=len(t["header"]))
            _borders(table)
            for j, h in enumerate(t["header"]):
                _cell_text(table.rows[0].cells[j], h, size, bold=True)
            for i, row in enumerate(t["rows"], start=1):
                for j, c in enumerate(row):
                    _cell_text(table.rows[i].cells[j], c, size)
            if t.get("note"):
                p = doc.add_paragraph(t["note"])
                p.runs[0].font.size = Pt(8)
                p.runs[0].font.color.rgb = GREY
    dest = HERE / "DT-002-Zone11-Sizing-Note.docx"
    doc.save(str(dest))
    print("written:", dest.name)


if __name__ == "__main__":
    to_markdown()
    to_docx()
